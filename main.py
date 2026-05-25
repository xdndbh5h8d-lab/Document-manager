import sys
from pathlib import Path

from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import (
    QApplication,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QListWidget,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


class DocumentManagerWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Document Manager")
        self.resize(900, 640)

        self.current_file_path = None
        self.current_file_type = None

        self._setup_ui()

    def _setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)

        self.workspace_list = QListWidget()
        self.workspace_list.setFixedWidth(220)
        self.workspace_list.itemClicked.connect(self._on_workspace_item_clicked)

        self.title_label = QLabel("Kein Dokument geladen")
        self.title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.content_editor = QTextEdit()

        self.file_path_input = QLineEdit()
        self.file_path_input.setReadOnly(True)

        open_docx_button = QPushButton("Word öffnen")
        open_docx_button.clicked.connect(self.open_word_document)

        open_pdf_button = QPushButton("PDF öffnen")
        open_pdf_button.clicked.connect(self.open_pdf_document)

        new_docx_button = QPushButton("Neue Word-Datei")
        new_docx_button.clicked.connect(self.create_new_word_document)

        new_pdf_button = QPushButton("Neue PDF-Datei")
        new_pdf_button.clicked.connect(self.create_new_pdf_document)

        save_button = QPushButton("Speichern")
        save_button.clicked.connect(self.save_document)

        button_layout = QHBoxLayout()
        button_layout.addWidget(open_docx_button)
        button_layout.addWidget(open_pdf_button)
        button_layout.addWidget(new_docx_button)
        button_layout.addWidget(new_pdf_button)
        button_layout.addWidget(save_button)

        right_layout = QVBoxLayout()
        right_layout.addWidget(self.title_label)
        right_layout.addWidget(self.file_path_input)
        right_layout.addLayout(button_layout)
        right_layout.addWidget(self.content_editor)

        main_layout = QHBoxLayout(central)
        main_layout.addWidget(self.workspace_list)
        main_layout.addLayout(right_layout)

        self._populate_default_workspace()

    def _populate_default_workspace(self):
        self.workspace_list.clear()
        self.workspace_list.addItem("Beispiel-Dokument 1")
        self.workspace_list.addItem("Beispiel-Dokument 2")
        self.workspace_list.addItem("Beispiel-Dokument 3")

    def _on_workspace_item_clicked(self, item):
        self.title_label.setText(f"Arbeitsbereich: {item.text()}")
        self.content_editor.setPlainText(
            "Hier kannst du Text eingeben und Dokumente verwalten."
        )
        self.current_file_path = None
        self.current_file_type = None
        self.file_path_input.clear()

    def open_word_document(self):
        path, _ = QFileDialog.getOpenFileName(self, "Word-Datei öffnen", "", "Word Dateien (*.docx)")
        if not path:
            return

        try:
            doc = Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.content_editor.setPlainText(text)
            self.title_label.setText(f"Word-Dokument: {Path(path).name}")
            self.current_file_path = Path(path)
            self.current_file_type = "docx"
            self.file_path_input.setText(path)
        except Exception as exc:
            QMessageBox.critical(self, "Fehler", f"Word-Datei konnte nicht geladen werden:\n{exc}")

    def open_pdf_document(self):
        path, _ = QFileDialog.getOpenFileName(self, "PDF-Datei öffnen", "", "PDF Dateien (*.pdf)")
        if not path:
            return

        try:
            reader = PdfReader(path)
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            self.content_editor.setPlainText("\n\n".join(texts))
            self.title_label.setText(f"PDF-Dokument: {Path(path).name}")
            self.current_file_path = Path(path)
            self.current_file_type = "pdf"
            self.file_path_input.setText(path)
        except Exception as exc:
            QMessageBox.critical(self, "Fehler", f"PDF-Datei konnte nicht geladen werden:\n{exc}")

    def create_new_word_document(self):
        self.current_file_path = None
        self.current_file_type = "docx"
        self.title_label.setText("Neue Word-Datei")
        self.content_editor.setPlainText("Neue Word-Datei\n")
        self.file_path_input.clear()

    def create_new_pdf_document(self):
        self.current_file_path = None
        self.current_file_type = "pdf"
        self.title_label.setText("Neue PDF-Datei")
        self.content_editor.setPlainText("Neue PDF-Datei\n")
        self.file_path_input.clear()

    def save_document(self):
        content = self.content_editor.toPlainText()
        if not content.strip():
            QMessageBox.warning(self, "Warnung", "Der Textbereich ist leer. Bitte füge Inhalt hinzu.")
            return

        if self.current_file_type == "docx":
            self._save_word(content)
        elif self.current_file_type == "pdf":
            self._save_pdf(content)
        else:
            choice = QMessageBox.question(
                self,
                "Dateityp wählen",
                "Möchtest du das Dokument als Word oder PDF speichern?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes,
            )
            if choice == QMessageBox.StandardButton.Yes:
                self.current_file_type = "docx"
                self._save_word(content)
            else:
                self.current_file_type = "pdf"
                self._save_pdf(content)

    def _save_word(self, content: str):
        path = self.current_file_path
        if not path or path.suffix.lower() != ".docx":
            path, _ = QFileDialog.getSaveFileName(self, "Word-Datei speichern", "", "Word Dateien (*.docx)")
            if not path:
                return
            path = Path(path if path.endswith(".docx") else f"{path}.docx")

        try:
            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)
            doc.save(path)
            self.current_file_path = path
            self.file_path_input.setText(str(path))
            self.title_label.setText(f"Word-Dokument: {path.name}")
            QMessageBox.information(self, "Erfolg", f"Word-Datei gespeichert:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Fehler", f"Word-Datei konnte nicht gespeichert werden:\n{exc}")

    def _save_pdf(self, content: str):
        path = self.current_file_path
        if not path or path.suffix.lower() != ".pdf":
            path, _ = QFileDialog.getSaveFileName(self, "PDF-Datei speichern", "", "PDF Dateien (*.pdf)")
            if not path:
                return
            path = Path(path if path.endswith(".pdf") else f"{path}.pdf")

        try:
            pdf = canvas.Canvas(str(path), pagesize=A4)
            text_object = pdf.beginText(40, A4[1] - 40)
            text_object.setFont("Helvetica", 11)
            for line in content.splitlines():
                if text_object.getY() < 40:
                    pdf.drawText(text_object)
                    pdf.showPage()
                    text_object = pdf.beginText(40, A4[1] - 40)
                    text_object.setFont("Helvetica", 11)
                text_object.textLine(line)
            pdf.drawText(text_object)
            pdf.save()
            self.current_file_path = path
            self.file_path_input.setText(str(path))
            self.title_label.setText(f"PDF-Dokument: {path.name}")
            QMessageBox.information(self, "Erfolg", f"PDF-Datei gespeichert:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Fehler", f"PDF-Datei konnte nicht gespeichert werden:\n{exc}")


def main():
    app = QApplication(sys.argv)
    window = DocumentManagerWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
