from pathlib import Path

from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def load_word(path: str) -> str:
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def load_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def save_word(content: str, path: str) -> None:
    doc = Document()
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(path)


def save_pdf(content: str, path: str) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    text_object = pdf.beginText(40, A4[1] - 40)
    text_object.setFont("Helvetica", 11)
    for line in content.splitlines():
        if text_object.getY() < 40:
            pdf.drawText(text_object)
            pdf.showPage()
            text_object = pdf.beginText(40, A4[1] - 40)
            text_object.setFont("Helvetica", 11)
        text_object.textLine(line)
    pdf.drawText(text_object)
    pdf.save()
